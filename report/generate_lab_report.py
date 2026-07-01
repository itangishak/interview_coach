# -*- coding: utf-8 -*-
"""Generate AI Interview Coach Lab Report as DOCX"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime
import os


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml('<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_header_row(table, headers, shade_color="1F2937"):
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, shade_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)


def add_table_data(table, data, start_row=1):
    for row_idx, row_data in enumerate(data):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + start_row].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"


def styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return h


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # TITLE PAGE
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI INTERVIEW COACH")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Real-Time Body Language and Behavioral Analysis System")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Lab Report")
    run.font.size = Pt(22)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Application-Oriented Course Design")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    today = datetime.date.today().strftime('%B %d, %Y')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Author: [Author Name]\nDate: " + today + "\n\nInstitution: [Institution Name]\nCourse: [Course Name]")
    run.font.size = Pt(12)
    
    doc.add_page_break()

    # TABLE OF CONTENTS
    p = doc.add_paragraph()
    run = p.add_run("TABLE OF CONTENTS")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    
    toc = [
        "1. Introduction",
        "   1.1 Purpose of the Lab Report",
        "   1.2 Project Overview",
        "   1.3 Key Contributions",
        "2. Functionality of the Project",
        "   2.1 Real-Time Body Tracking",
        "   2.2 Pose Estimation",
        "   2.3 Pose Visualization",
        "   2.4 Posture Assessment",
        "   2.5 Session Reporting",
        "3. Implementation Ideas",
        "   3.1 Computer Vision Pipeline",
        "   3.2 Scale-Independent Normalization",
        "   3.3 3-D Head Pose Estimation",
        "   3.4 Temporal Smoothing with EMA",
        "   3.5 Detection Gating",
        "   3.6 Per-User Calibration",
        "   3.7 Smile Detection Improvements",
        "   3.8 Software Engineering Design Patterns",
        "   3.9 WebSocket Real-Time Communication",
        "4. Technology Stack",
        "5. Development Environment",
        "6. User Interaction",
        "7. Testing",
        "8. Critical Bug Fix: WebSocket Deadlock",
        "9. Conclusion",
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        if not item.startswith("   "):
            run.font.bold = True
        run.font.size = Pt(11)
    
    doc.add_page_break()

    # SECTION 1: INTRODUCTION
    styled_heading(doc, "1. Introduction", 1)
    styled_heading(doc, "1.1 Purpose of the Lab Report", 2)
    doc.add_paragraph(
        "The purpose of this lab report is to document the application-oriented course design "
        "for the AI Interview Coach. The project aims to develop an application that can track "
        "and analyze body language, eye contact, facial expressions, and movement in real-time "
        "via webcam, then deliver live feedback through a browser dashboard. This report will "
        "outline the functionality of the project, implementation ideas, technology stack, "
        "development environment, user interaction, testing approach, and a critical WebSocket "
        "deadlock bug that was identified and resolved during development."
    )
    
    styled_heading(doc, "1.2 Project Overview", 2)
    doc.add_paragraph(
        "The AI Interview Coach is a real-time interview coaching system that analyzes five "
        "behavioral metrics: eye contact, smile, posture, head stability, and body movement. "
        "The system processes video frames at approximately 15 fps, extracts 478 facial landmarks "
        "using MediaPipe Face Mesh and 33 body landmarks using MediaPipe Pose, computes behavioral "
        "metrics using scale-independent geometric formulas, and presents results through an "
        "interactive Next.js web dashboard with real-time WebSocket communication. "
        "A persistent SQLite database stores session data and per-user calibration profiles."
    )
    
    styled_heading(doc, "1.3 Key Contributions", 2)
    for item in [
        "Scale-independent metric computation normalized by interocular distance and shoulder width.",
        "3-D head pose estimation via OpenCV solvePnP with camera-above-screen offset correction.",
        "Per-user calibration with persistent cross-session profiles stored in SQLite.",
        "Real-time bidirectional WebSocket communication between frontend and FastAPI backend.",
        "15+ software engineering design patterns verified by 121+ dedicated unit tests.",
        "Critical threading.Lock deadlock resolution preventing WebSocket handshake failure.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # SECTION 2: FUNCTIONALITY
    styled_heading(doc, "2. Functionality of the Project", 1)
    styled_heading(doc, "2.1 Real-Time Body Tracking", 2)
    doc.add_paragraph(
        "The application uses computer vision to track body movements via webcam at 15 fps. "
        "It tracks 478 facial landmarks (MediaPipe Face Mesh) and 33 body landmarks (MediaPipe Pose), "
        "plus 6-DOF head orientation estimated via Perspective-n-Point solving."
    )
    
    styled_heading(doc, "2.2 Pose Estimation", 2)
    doc.add_paragraph("Five behavioral metrics are computed with the following weights:")
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["Metric", "Weight", "Source"])
    add_table_data(table, [
        ["Eye Contact", "0.30", "Iris position relative to eye-socket width (ICD-normalized)"],
        ["Posture", "0.20", "Shoulder tilt + lean penalty + hip angle"],
        ["Head Stability", "0.20", "Nose-position std over rolling window (ICD-normalized)"],
        ["Body Movement", "0.15", "Shoulder + nose variance (shoulder-width-normalized)"],
        ["Smile", "0.15", "Mouth ratio + lip elevation + cheek-squint proxy"],
    ])
    doc.add_paragraph("")
    
    styled_heading(doc, "2.3 Pose Visualization", 2)
    doc.add_paragraph(
        "The dashboard displays live metric cards, a confidence ring (0-100%), sparkline charts, "
        "contextual feedback, head orientation status, and optional diagnostic overlays."
    )
    
    styled_heading(doc, "2.4 Posture Assessment", 2)
    doc.add_paragraph(
        "Posture assessment includes shoulder tilt analysis, lean detection, spine angle estimation, "
        "and per-user baseline comparison for personalized assessment."
    )
    
    styled_heading(doc, "2.5 Session Reporting", 2)
    doc.add_paragraph(
        "At session end, a comprehensive summary is generated with duration, frame counts, "
        "per-metric statistics, percentile confidence scores, and recommendations."
    )

    # SECTION 3: IMPLEMENTATION IDEAS
    styled_heading(doc, "3. Implementation Ideas", 1)
    styled_heading(doc, "3.1 Computer Vision Pipeline", 2)
    doc.add_paragraph(
        "Seven-stage pipeline: frame capture, MediaPipe inference, 3-D head pose estimation via "
        "solvePnP, metric computation with scale-independent normalization, EMA temporal smoothing, "
        "confidence scoring, and feedback generation."
    )
    
    styled_heading(doc, "3.2 Scale-Independent Normalization", 2)
    doc.add_paragraph(
        "All distances are normalized by interocular distance (facial) or shoulder width (body) "
        "to eliminate camera-distance dependence."
    )
    
    styled_heading(doc, "3.3 3-D Head Pose Estimation", 2)
    doc.add_paragraph(
        "cv2.solvePnP with 6 landmarks maps to a canonical 3-D face model. Euler angles (yaw/pitch/roll) "
        "are extracted. Camera offset correction ensures genuine eye contact reads as 0 degrees."
    )
    
    styled_heading(doc, "3.4 Temporal Smoothing with EMA", 2)
    doc.add_paragraph(
        "Exponential moving average (alpha=0.30 for most metrics, 0.45 for smile) replaces the "
        "simple mean. The 80th percentile of a 30-frame window captures smile events."
    )
    
    styled_heading(doc, "3.5 Detection Gating", 2)
    doc.add_paragraph(
        "When no face/pose is detected, the system freezes EMA state, emits 0.0 wire values, "
        "sets validity flags to False, and excludes frames from aggregates."
    )
    
    styled_heading(doc, "3.6 Per-User Calibration", 2)
    doc.add_paragraph(
        "Two-phase calibration: 2-second in-session neutral baseline, plus persistent SQLite "
        "profiles with EMA blending (30% new, 70% stored history)."
    )
    
    styled_heading(doc, "3.7 Smile Detection Improvements", 2)
    doc.add_paragraph(
        "Three signals: ratio score (suppressed when jaw open), elevation score, and cheek-squint "
        "proxy. Blend: 70% mouth geometry, 30% cheek squint."
    )
    
    styled_heading(doc, "3.8 Software Engineering Design Patterns", 2)
    doc.add_paragraph(
        "The backend implements 15+ design patterns verified by 121+ unit tests across test_patterns.py "
        "and test_interview_metrics.py:"
    )
    
    table = doc.add_table(rows=14, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ["Pattern", "File", "Key Element"])
    add_table_data(table, [
        ["Singleton + Metaclass", "core/singleton.py", "SingletonMeta (lock-free)"],
        ["Decorator Factory", "core/decorators.py", "@validate_score, @log_call, @retry, @memoize"],
        ["Template Method", "services/metric_scorers.py", "BaseMetricScorer.update()"],
        ["Factory Method", "services/metric_scorers.py", "BaseMetricScorer.from_config()"],
        ["Strategy", "services/metric_scorers.py", "5 scorers with different score()"],
        ["Abstract Base Class", "services/metric_scorers.py", "BaseMetricScorer(ABC)"],
        ["Iterator Protocol", "core/iterators.py", "FrameWindowIterator, SentenceTokenIterator"],
        ["Generator", "core/iterators.py", "metric_history_generator, frame_chunk_generator"],
        ["Context Manager", "services/session_service.py", "session_scope()"],
        ["Dataclass + Slots", "services/metric_scorers.py", "MetricResult(slots=True)"],
        ["Closure", "services/metric_scorers.py", "make_ema_fn, make_threshold_checker"],
        ["Property / Classmethod", "services/metric_scorers.py", "@property, @classmethod, @staticmethod"],
        ["Inheritance + Polymorphism", "services/feedback_engine.py", "BaseFeedbackEngine hierarchy"],
    ])
    doc.add_paragraph("")

    styled_heading(doc, "3.9 WebSocket Real-Time Communication", 2)
    doc.add_paragraph(
        "The frontend establishes a persistent WebSocket connection to the FastAPI backend "
        "via two endpoints: /ws/interview for interview coaching (frame analysis) and /ws/sign "
        "for sign-language prediction. The interview WebSocket receives base64-encoded webcam frames, "
        "processes them through the InterviewAnalyzer pipeline, and streams per-frame analysis "
        "results back to the client. All blocking database operations are delegated to a thread "
        "pool via asyncio.run_in_executor to keep the event loop responsive."
    )

    # SECTION 4: TECHNOLOGY STACK
    styled_heading(doc, "4. Technology Stack", 1)
    styled_heading(doc, "4.1 Programming Language", 2)
    for item in [
        "Python 3.10+: Backend language (FastAPI, OpenCV, MediaPipe)",
        "TypeScript/JavaScript: Frontend for Next.js dashboard",
        "HTML5/CSS3: Web interface markup and styling",
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    styled_heading(doc, "4.2 Libraries", 2)
    doc.add_paragraph("Backend:")
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, ["Library", "Version", "Purpose"])
    add_table_data(table, [
        ["FastAPI", "0.111+", "Web framework"],
        ["Uvicorn", "0.30+", "ASGI server"],
        ["OpenCV", "4.10+", "Video capture, solvePnP"],
        ["MediaPipe", "0.10+", "Face Mesh + Pose"],
        ["SQLAlchemy", "2.0+", "ORM"],
        ["NumPy", "1.26+", "Numerical computations"],
        ["Pydantic", "2.7+", "Data validation"],
        ["Pillow", "10.3+", "Image encoding"],
    ])
    doc.add_paragraph("")
    doc.add_paragraph("Frontend:")
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, ["Library", "Version", "Purpose"])
    add_table_data(table, [
        ["Next.js", "14.2.5", "React framework"],
        ["React", "18.3.1", "UI library"],
        ["Tailwind CSS", "3.4.6", "CSS framework"],
        ["WebSocket API", "Native", "Real-time communication"],
    ])
    doc.add_paragraph("")

    # SECTION 5: DEVELOPMENT ENVIRONMENT
    styled_heading(doc, "5. Development Environment", 1)
    doc.add_paragraph("Prerequisites: Python 3.10+, Node.js 18+, Git, webcam.")
    doc.add_paragraph("Backend setup: cd backend/ -> python -m venv .venv -> source .venv/bin/activate -> pip install -r requirements.txt")
    doc.add_paragraph("Frontend setup: cd frontend/ -> npm install")
    doc.add_paragraph("Run: Backend (uvicorn app.main:app --reload on :8000) + Frontend (npm run dev on :3000)")

    # SECTION 6: USER INTERACTION
    styled_heading(doc, "6. User Interaction", 1)
    doc.add_paragraph(
        "The GUI displays webcam feed with pose overlays. Users start sessions, undergo 2-second "
        "calibration, receive live metrics and feedback, view session reports, and access diagnostic "
        "mode with geometric overlays and raw data readouts."
    )
    doc.add_paragraph("Workflow: Start -> Calibrate -> Live Tracking -> Feedback -> Stop -> Report")

    # SECTION 7: TESTING
    styled_heading(doc, "7. Testing", 1)
    doc.add_paragraph("The test suite comprises 121+ tests across 2 files:")
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    add_header_row(table, ["Test File", "Test Methods", "Coverage"])
    add_table_data(table, [
        ["test_interview_metrics.py", "46", "Metric functions, EMA, hysteresis, gating, solvePnP, user profiles"],
        ["test_patterns.py", "75+", "Design patterns: Singleton, decorators, closures, iterators, generators, OOP hierarchy"],
    ])
    doc.add_paragraph("")
    doc.add_paragraph(
        "Integration testing covers WebSocket pipeline, database operations, API endpoints, "
        "and frontend-backend data flow. User acceptance testing evaluated usability, performance, "
        "and accuracy with real users."
    )

    # SECTION 8: CRITICAL BUG FIX
    styled_heading(doc, "8. Critical Bug Fix: WebSocket Deadlock", 1)
    styled_heading(doc, "8.1 Symptom", 2)
    doc.add_paragraph(
        "The Start button on the frontend remained permanently greyed out (disabled). The button "
        "state was controlled by a connected flag that only became true when the WebSocket onopen "
        "event fired. The onopen event never fired because the WebSocket handshake never completed."
    )
    
    styled_heading(doc, "8.2 Root Cause Analysis", 2)
    doc.add_paragraph(
        "The root cause was identified as a threading.Lock deadlock inside the asyncio event loop. "
        "The chain of events was:"
    )
    for item in [
        "WebSocket handler (asyncio coroutine, event loop thread) calls SessionService()",
        "SessionService() triggers first-time DatabaseManager() creation via SingletonMeta.__call__",
        "SingletonMeta.__call__ acquires threading.Lock (_SINGLETON_LOCK)",
        "Inside the lock, DatabaseManager.__init__ runs synchronous SQLite I/O (create_engine + Base.metadata.create_all)",
        "The asyncio event loop thread is completely blocked holding the lock",
        "The event loop cannot process the WebSocket upgrade handshake response",
        "The browser waits indefinitely for the 101 Switching Protocols response",
        "ws.onopen never fires, connected stays false, the Start button stays disabled",
    ]:
        doc.add_paragraph(item, style='List Number')
    
    styled_heading(doc, "8.3 Confirmation", 2)
    doc.add_paragraph(
        "The deadlock was confirmed via system-level inspection: /proc/<pid>/wchan showed "
        "hrtimer_nanosleep for the watcher (alive but sleeping) while the worker was stuck on "
        "futex_do_wait — the futex being the internal mechanism of threading.Lock."
    )
    
    styled_heading(doc, "8.4 Resolution", 2)
    doc.add_paragraph("Two complementary fixes were applied:")
    
    doc.add_paragraph(
        "Fix 1 — Lock-Free SingletonMeta: The threading.Lock was removed from SingletonMeta entirely. "
        "It was replaced with a lock-free instance-caching pattern. Python's Global Interpreter Lock (GIL) "
        "makes dict get/set operations atomic enough for this pattern. In the worst-case race condition, "
        "one extra instance is created and immediately discarded — no thread ever blocks, so the asyncio "
        "event loop cannot freeze."
    )
    
    doc.add_paragraph(
        "Fix 2 — FastAPI Lifespan Pre-Warming: A lifespan context manager was added to main.py that "
        "pre-initializes SessionService and UserProfileService in a thread pool (via run_in_executor) "
        "before the application starts serving traffic. This ensures the singletons are already "
        "instantiated when the first WebSocket client connects, making the fast path in SingletonMeta "
        "always win in practice."
    )
    
    styled_heading(doc, "8.5 Why the Fix Works", 2)
    for item in [
        "The lock-free singleton means calling SessionService() inside an async handler cannot freeze the event loop.",
        "The lifespan pre-warming ensures the heavy SQLite I/O happens before any client connects.",
        "Even if a race occurs (e.g., during testing), the extra instance is discarded and the event loop continues.",
        "The fix eliminates both the primary deadlock and the secondary circular-import re-acquisition risk.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # SECTION 9: CONCLUSION
    styled_heading(doc, "9. Conclusion", 1)
    doc.add_paragraph(
        "The AI Interview Coach successfully demonstrates real-time behavioral analysis combining "
        "computer vision, geometric normalization, and modern web technologies. The system achieves "
        "real-time tracking at 15 fps with 5 behavioral metrics, scale-independent normalization, "
        "per-user calibration, 15+ design patterns verified by 121+ tests, and honest detection gating."
    )
    doc.add_paragraph(
        "The critical WebSocket deadlock bug — caused by a threading.Lock acquired on the asyncio "
        "event loop thread during synchronous SQLite I/O — was resolved through a lock-free singleton "
        "redesign and FastAPI lifespan pre-warming. This case study underscores the importance of "
        "understanding the interaction between synchronous primitives (threading.Lock) and asynchronous "
        "event loops in Python web applications."
    )
    doc.add_paragraph(
        "The technology stack provides a solid foundation for future extensibility, with "
        "configuration-driven tuning requiring no code changes."
    )

    # Save
    output_path = os.path.join(script_dir, "AI_Interview_Coach_Lab_Report.docx")
    doc.save(output_path)
    
    # Verify
    with open(os.path.join(script_dir, "status.txt"), "w") as f:
        f.write("File created: " + output_path + "\n")
        f.write("Size: " + str(os.path.getsize(output_path)) + " bytes\n")
    
    print(f"Report generated successfully: {output_path}")
    print(f"File size: {os.path.getsize(output_path)} bytes")


if __name__ == "__main__":
    main()